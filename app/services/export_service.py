"""
Serviço para exportação de revisões em PDF/DOCX
"""

import os
import logging
from typing import Dict
from flask import render_template_string

logger = logging.getLogger(__name__)


class ExportService:
    """Serviço para exportar revisões em diferentes formatos"""
    
    def export_to_pdf(self, review_data: Dict) -> bytes:
        """Exporta revisão para PDF"""
        try:
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
            from reportlab.lib import colors
            from io import BytesIO
            
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=A4)
            story = []
            styles = getSampleStyleSheet()
            
            # Título
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                textColor=colors.HexColor('#8B5CF6'),
                spaceAfter=30
            )
            story.append(Paragraph(f"Revisão Jurídica - {review_data.get('title', 'Documento')}", title_style))
            story.append(Spacer(1, 0.2*inch))
            
            # Informações do documento
            story.append(Paragraph("<b>Informações do Documento</b>", styles['Heading2']))
            story.append(Paragraph(f"<b>Título:</b> {review_data.get('title', 'N/A')}", styles['Normal']))
            story.append(Paragraph(f"<b>Resumo:</b> {review_data.get('summary', 'N/A')}", styles['Normal']))
            story.append(Paragraph(f"<b>Descrição:</b> {review_data.get('description', 'N/A')}", styles['Normal']))
            story.append(Spacer(1, 0.2*inch))
            
            # Informações da revisão
            story.append(Paragraph("<b>Informações da Revisão</b>", styles['Heading2']))
            story.append(Paragraph(f"<b>Versão:</b> {review_data.get('version', 'N/A')}", styles['Normal']))
            story.append(Paragraph(f"<b>Revisor:</b> {review_data.get('reviewer_name', 'N/A')}", styles['Normal']))
            story.append(Paragraph(f"<b>Data:</b> {review_data.get('review_date', 'N/A')}", styles['Normal']))
            story.append(Paragraph(f"<b>Comentários:</b> {review_data.get('comments', 'N/A')}", styles['Normal']))
            story.append(Spacer(1, 0.2*inch))
            
            # Riscos
            if review_data.get('risks'):
                story.append(Paragraph("<b>Riscos Identificados</b>", styles['Heading2']))
                for risk in review_data.get('risks', []):
                    story.append(Paragraph(f"<b>Risco:</b> {risk.get('risk_text', 'N/A')}", styles['Normal']))
                    story.append(Paragraph(f"<b>Sugestão:</b> {risk.get('legal_suggestion', 'N/A')}", styles['Normal']))
                    story.append(Paragraph(f"<b>Definição Final:</b> {risk.get('final_definition', 'N/A')}", styles['Normal']))
                    story.append(Spacer(1, 0.1*inch))
            
            # Observações
            if review_data.get('observations'):
                story.append(Paragraph("<b>Observações Gerais</b>", styles['Heading2']))
                story.append(Paragraph(review_data.get('observations', ''), styles['Normal']))
                story.append(Spacer(1, 0.2*inch))
            
            # Aprovações
            if review_data.get('approvals'):
                story.append(Paragraph("<b>Histórico de Aprovações</b>", styles['Heading2']))
                for approval in review_data.get('approvals', []):
                    story.append(Paragraph(f"<b>Aprovador:</b> {approval.get('approver_name', 'N/A')}", styles['Normal']))
                    story.append(Paragraph(f"<b>Status:</b> {approval.get('status', 'N/A')}", styles['Normal']))
                    approved_at = approval.get('approved_at', 'N/A')
                    if approved_at and approved_at != 'N/A' and hasattr(approved_at, 'strftime'):
                        approved_at = approved_at.strftime('%d/%m/%Y %H:%M:%S')
                    story.append(Paragraph(f"<b>Data:</b> {approved_at}", styles['Normal']))
                    story.append(Paragraph(f"<b>Comentário:</b> {approval.get('comments', 'N/A')}", styles['Normal']))
                    story.append(Spacer(1, 0.1*inch))
            
            doc.build(story)
            buffer.seek(0)
            return buffer.getvalue()
        except Exception as e:
            logger.error(f"Erro ao exportar para PDF: {str(e)}")
            raise
    
    def export_to_pdf_with_history(self, review_data: Dict, versions_with_comments: list, 
                                   versions_with_risks: list) -> bytes:
        """Exporta revisão para PDF incluindo histórico completo"""
        try:
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
            from reportlab.lib import colors
            from io import BytesIO
            
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=A4)
            story = []
            styles = getSampleStyleSheet()
            
            # Título
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                textColor=colors.HexColor('#8B5CF6'),
                spaceAfter=30
            )
            story.append(Paragraph(f"Revisão Jurídica - {review_data.get('title', 'Documento')}", title_style))
            story.append(Paragraph("<i>Histórico Completo</i>", styles['Normal']))
            story.append(Spacer(1, 0.2*inch))
            
            # Informações do documento
            story.append(Paragraph("<b>Informações do Documento</b>", styles['Heading2']))
            story.append(Paragraph(f"<b>Título:</b> {review_data.get('title', 'N/A')}", styles['Normal']))
            story.append(Paragraph(f"<b>Descrição:</b> {review_data.get('description', 'N/A')}", styles['Normal']))
            story.append(Paragraph(f"<b>Versão Atual:</b> v{review_data.get('version', 'N/A')}", styles['Normal']))
            story.append(Spacer(1, 0.3*inch))
            
            # Histórico de Revisões
            if versions_with_comments:
                story.append(PageBreak())
                story.append(Paragraph("<b>Histórico de Revisões</b>", styles['Heading2']))
                story.append(Spacer(1, 0.1*inch))
                
                for version in versions_with_comments:
                    story.append(Paragraph(f"<b>Versão {version.get('version', 'N/A')}</b>", styles['Heading3']))
                    story.append(Paragraph(f"<b>Responsável:</b> {version.get('reviewer_name', 'N/A')}", styles['Normal']))
                    
                    review_date = version.get('review_date', 'N/A')
                    if review_date and review_date != 'N/A' and hasattr(review_date, 'strftime'):
                        review_date = review_date.strftime('%d/%m/%Y %H:%M:%S')
                    story.append(Paragraph(f"<b>Data/Hora:</b> {review_date}", styles['Normal']))
                    
                    comments_list = version.get('comments_list', [])
                    if comments_list:
                        story.append(Paragraph("<b>Comentários:</b>", styles['Normal']))
                        for comment in comments_list:
                            comment_date = comment.get('review_date', 'N/A')
                            if comment_date and comment_date != 'N/A' and hasattr(comment_date, 'strftime'):
                                comment_date = comment_date.strftime('%d/%m/%Y %H:%M:%S')
                            story.append(Paragraph(
                                f"• <i>{comment.get('reviewer_name', 'N/A')} - {comment_date}</i>", 
                                styles['Normal']
                            ))
                            story.append(Paragraph(f"  {comment.get('comment', 'N/A')}", styles['Normal']))
                    else:
                        story.append(Paragraph("<i>Nenhum comentário</i>", styles['Normal']))
                    
                    story.append(Spacer(1, 0.2*inch))
            
            # Histórico de Riscos
            if versions_with_risks:
                story.append(PageBreak())
                story.append(Paragraph("<b>Histórico de Riscos</b>", styles['Heading2']))
                story.append(Spacer(1, 0.1*inch))
                
                for version in versions_with_risks:
                    story.append(Paragraph(f"<b>Versão {version.get('version', 'N/A')}</b>", styles['Heading3']))
                    
                    risks_list = version.get('risks_list', [])
                    if risks_list:
                        for risk in risks_list:
                            story.append(Paragraph(f"<b>Risco:</b> {risk.get('risk_text', 'N/A')}", styles['Normal']))
                            if risk.get('legal_suggestion'):
                                story.append(Paragraph(f"<b>Sugestão do Jurídico:</b> {risk.get('legal_suggestion', '')}", styles['Normal']))
                            if risk.get('final_definition'):
                                story.append(Paragraph(f"<b>Definição Final:</b> {risk.get('final_definition', '')}", styles['Normal']))
                            story.append(Spacer(1, 0.1*inch))
                    else:
                        story.append(Paragraph("<i>Nenhum risco identificado</i>", styles['Normal']))
                    
                    story.append(Spacer(1, 0.2*inch))
            
            # Observações da versão atual
            if review_data.get('observations'):
                story.append(PageBreak())
                story.append(Paragraph("<b>Observações Gerais (Versão Atual)</b>", styles['Heading2']))
                story.append(Paragraph(review_data.get('observations', ''), styles['Normal']))
            
            doc.build(story)
            buffer.seek(0)
            return buffer.getvalue()
        except Exception as e:
            logger.error(f"Erro ao exportar para PDF com histórico: {str(e)}")
            raise
    
    def export_to_docx(self, review_data: Dict) -> bytes:
        """Exporta revisão para DOCX"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from io import BytesIO
            
            doc = Document()
            
            # Título
            title = doc.add_heading(f"Revisão Jurídica - {review_data.get('title', 'Documento')}", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Informações do documento
            doc.add_heading('Informações do Documento', 1)
            doc.add_paragraph(f"Título: {review_data.get('title', 'N/A')}")
            doc.add_paragraph(f"Resumo: {review_data.get('summary', 'N/A')}")
            doc.add_paragraph(f"Descrição: {review_data.get('description', 'N/A')}")
            
            # Informações da revisão
            doc.add_heading('Informações da Revisão', 1)
            doc.add_paragraph(f"Versão: {review_data.get('version', 'N/A')}")
            doc.add_paragraph(f"Revisor: {review_data.get('reviewer_name', 'N/A')}")
            review_date = review_data.get('review_date', 'N/A')
            if isinstance(review_date, str):
                review_date = review_date
            elif hasattr(review_date, 'strftime'):
                review_date = review_date.strftime('%d/%m/%Y %H:%M:%S')
            doc.add_paragraph(f"Data: {review_date}")
            doc.add_paragraph(f"Comentários: {review_data.get('comments', 'N/A')}")
            
            # Riscos
            if review_data.get('risks'):
                doc.add_heading('Riscos Identificados', 1)
                for risk in review_data.get('risks', []):
                    doc.add_paragraph(f"Risco: {risk.get('risk_text', 'N/A')}", style='List Bullet')
                    doc.add_paragraph(f"Sugestão: {risk.get('legal_suggestion', 'N/A')}")
                    doc.add_paragraph(f"Definição Final: {risk.get('final_definition', 'N/A')}")
            
            # Observações
            if review_data.get('observations'):
                doc.add_heading('Observações Gerais', 1)
                doc.add_paragraph(review_data.get('observations', ''))
            
            # Aprovações
            if review_data.get('approvals'):
                doc.add_heading('Histórico de Aprovações', 1)
                for approval in review_data.get('approvals', []):
                    doc.add_paragraph(f"Aprovador: {approval.get('approver_name', 'N/A')}", style='List Bullet')
                    doc.add_paragraph(f"Status: {approval.get('status', 'N/A')}")
                    approved_at = approval.get('approved_at', 'N/A')
                    if approved_at and approved_at != 'N/A' and hasattr(approved_at, 'strftime'):
                        approved_at = approved_at.strftime('%d/%m/%Y %H:%M:%S')
                    doc.add_paragraph(f"Data: {approved_at}")
                    doc.add_paragraph(f"Comentário: {approval.get('comments', 'N/A')}")
            
            # Salvar em buffer
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()
        except Exception as e:
            logger.error(f"Erro ao exportar para DOCX: {str(e)}")
            raise
    
    def export_to_docx_with_history(self, review_data: Dict, versions_with_comments: list,
                                    versions_with_risks: list) -> bytes:
        """Exporta revisão para DOCX incluindo histórico completo"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from io import BytesIO
            
            doc = Document()
            
            # Título
            title = doc.add_heading(f"Revisão Jurídica - {review_data.get('title', 'Documento')}", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle = doc.add_paragraph("Histórico Completo")
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Informações do documento
            doc.add_heading('Informações do Documento', 1)
            doc.add_paragraph(f"Título: {review_data.get('title', 'N/A')}")
            doc.add_paragraph(f"Descrição: {review_data.get('description', 'N/A')}")
            doc.add_paragraph(f"Versão Atual: v{review_data.get('version', 'N/A')}")
            
            # Histórico de Revisões
            if versions_with_comments:
                doc.add_page_break()
                doc.add_heading('Histórico de Revisões', 1)
                
                for version in versions_with_comments:
                    doc.add_heading(f"Versão {version.get('version', 'N/A')}", 2)
                    doc.add_paragraph(f"Responsável: {version.get('reviewer_name', 'N/A')}")
                    
                    review_date = version.get('review_date', 'N/A')
                    if review_date and review_date != 'N/A' and hasattr(review_date, 'strftime'):
                        review_date = review_date.strftime('%d/%m/%Y %H:%M:%S')
                    doc.add_paragraph(f"Data/Hora: {review_date}")
                    
                    comments_list = version.get('comments_list', [])
                    if comments_list:
                        doc.add_paragraph("Comentários:", style='Heading 3')
                        for comment in comments_list:
                            comment_date = comment.get('review_date', 'N/A')
                            if comment_date and comment_date != 'N/A' and hasattr(comment_date, 'strftime'):
                                comment_date = comment_date.strftime('%d/%m/%Y %H:%M:%S')
                            doc.add_paragraph(
                                f"{comment.get('reviewer_name', 'N/A')} - {comment_date}",
                                style='List Bullet'
                            )
                            doc.add_paragraph(f"  {comment.get('comment', 'N/A')}")
                    else:
                        doc.add_paragraph("Nenhum comentário")
            
            # Histórico de Riscos
            if versions_with_risks:
                doc.add_page_break()
                doc.add_heading('Histórico de Riscos', 1)
                
                for version in versions_with_risks:
                    doc.add_heading(f"Versão {version.get('version', 'N/A')}", 2)
                    
                    risks_list = version.get('risks_list', [])
                    if risks_list:
                        for risk in risks_list:
                            doc.add_paragraph(f"Risco: {risk.get('risk_text', 'N/A')}", style='List Bullet')
                            if risk.get('legal_suggestion'):
                                doc.add_paragraph(f"Sugestão do Jurídico: {risk.get('legal_suggestion', '')}")
                            if risk.get('final_definition'):
                                doc.add_paragraph(f"Definição Final: {risk.get('final_definition', '')}")
                    else:
                        doc.add_paragraph("Nenhum risco identificado")
            
            # Observações da versão atual
            if review_data.get('observations'):
                doc.add_page_break()
                doc.add_heading('Observações Gerais (Versão Atual)', 1)
                doc.add_paragraph(review_data.get('observations', ''))
            
            # Salvar em buffer
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()
        except Exception as e:
            logger.error(f"Erro ao exportar para DOCX com histórico: {str(e)}")
            raise


# Instância global do serviço
export_service = ExportService()

